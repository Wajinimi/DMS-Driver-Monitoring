"""Surgically update the MSc report docx: embed real figures, add the training
dynamics and confusion-structure analysis, and trim prose to hold the word limit.

Run with --dry to print the resulting word count without writing the file.
"""

import copy
import shutil
import sys
import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

BASE = "/Users/dell/Library/CloudStorage/OneDrive-hull.ac.uk/MEETING"
DOCX = os.path.join(BASE, "MSc_Research_Project_Report_GAN_CXR.docx")
FIG = os.path.join(BASE, "report_figures")
WORKFLOW = os.path.join(BASE, "Experiment_Workflow_Diagram.png")

# ---------------------------------------------------------------- replacements
# Keyed by the paragraph index in the original document.

REPLACE = {
    3: (
        "Deep convolutional neural networks perform strongly on chest radiograph "
        "interpretation, yet their accuracy on infrequent thoracic findings remains "
        "constrained by severe class imbalance, for which generative adversarial "
        "networks (GANs) are widely promoted as a remedy. This study evaluates whether "
        "label-conditioned GAN augmentation improves multi-class chest radiograph "
        "classification relative to classical geometric and intensity-based alternatives. "
        "A seven-class subset of 85,116 frontal radiographs, 70.9% labelled No Finding, was "
        "partitioned by stratified 80/20 sampling into 68,092 training and 17,024 validation "
        "images. Three architectures, a scratch-trained custom convolutional network and "
        "ImageNet-pretrained ResNet50 and VGG19 with frozen bases, were each trained under no "
        "augmentation, classical geometric transformation, intensity perturbation and "
        "conditional GAN augmentation. The twelve experiments were compared on macro-averaged "
        "precision, recall and F1 using bootstrap resampling with 2,000 replicates at a "
        "significance level of 0.05. Conditional GAN augmentation failed to improve macro-F1 "
        "for any architecture and degraded it significantly in every case, precipitating a "
        "collapse to 0.038 against a baseline of 0.153 for the scratch-trained network. "
        "Intensity augmentation with VGG19 "
        "achieved the highest macro-F1 of 0.175, statistically indistinguishable from the "
        "ResNet50 baseline of 0.174, and no configuration exceeded the majority-class prior "
        "of 0.709 in accuracy. These findings indicate that an under-trained conditional GAN "
        "functions as label noise rather than as informative augmentation."
    ),
    6: (
        "Chest radiography remains the most frequently performed diagnostic imaging "
        "examination worldwide, and the interpretative burden it places on radiology "
        "services has motivated sustained interest in automated analysis. The release of "
        "large, weakly labelled archives such as ChestX-ray8 transformed this field by "
        "enabling convolutional neural networks to be trained at a scale previously "
        "confined to natural-image benchmarks (Wang et al., 2017). Subsequent systems, most "
        "prominently CheXNet, reported pneumonia detection approaching radiologist-level "
        "agreement, and CheXpert extended the paradigm with explicit uncertainty labelling "
        "(Rajpurkar et al., 2017; Irvin et al., 2019). "
        "One obstacle has nonetheless proved resistant to increases in raw data volume: the "
        "distribution of thoracic findings is profoundly long-tailed, so the abnormalities of "
        "greatest clinical consequence are those observed least often. Models optimised on "
        "such data learn the majority prior and generalise poorly to rare classes, a failure "
        "mode conventional accuracy reporting tends to conceal (Holste et al., 2022)."
    ),
    7: (
        "Data augmentation is the standard countermeasure. Classical geometric and "
        "photometric transformations expand the effective size of a training set at "
        "negligible cost and are established components of the deep learning pipeline "
        "(Shorten and Khoshgoftaar, 2019). Their expressive range is bounded, however, because "
        "every synthesised example remains an affine or intensity-scaled restatement of an "
        "existing image and introduces no new anatomical variation. Generative adversarial "
        "networks appear to escape this constraint by learning the data distribution directly "
        "and sampling from it (Goodfellow et al., 2014), while conditional variants steer "
        "generation towards a designated class (Mirza and Osindero, 2014). This has obvious "
        "appeal for imbalanced medical corpora, and a substantial literature reports accuracy "
        "gains from GAN-derived synthetic radiographs (Salehinejad et al., 2018; Waheed et al., 2020; "
        "Karbhari et al., 2021)."
    ),
    8: (
        "That literature is, however, methodologically narrow. The great majority of "
        "favourable reports concern binary or three-class problems, frequently involving "
        "COVID-19 cohorts of a few thousand images, and few evaluate GAN augmentation against "
        "a properly matched classical baseline under identical training conditions. Fewer "
        "still quantify whether observed differences exceed sampling variability. It therefore "
        "remains unclear whether the reported advantages generalise to a large, genuinely "
        "multi-class corpus in which the abundance of real data may already exhaust what "
        "synthetic images can contribute."
    ),
    9: (
        "This study addresses that question directly. Its aim is to determine whether "
        "conditional GAN augmentation confers a measurable and statistically defensible "
        "advantage over classical and intensity-based augmentation for seven-class chest "
        "radiograph classification. Four objectives follow: to construct a controlled "
        "factorial comparison in which augmentation regime is the sole manipulated variable "
        "across three architectures of contrasting inductive bias; to evaluate performance "
        "using macro-averaged metrics that weight rare classes equally with common ones; to "
        "establish the reliability of every comparison through bootstrap resampling rather "
        "than point estimates alone; and to interpret the resulting evidence against the "
        "wider literature, identifying the conditions under which synthetic augmentation is "
        "likely to help or to harm."
    ),
    11: (
        "Automated chest radiograph interpretation was reshaped by ChestX-ray8, which "
        "supplied 112,120 frontal images labelled for common thoracic pathologies through "
        "natural-language processing of radiology reports (Wang et al., 2017). The dataset "
        "established the benchmark on which most subsequent work is calibrated, but its "
        "construction embedded two difficulties that persist here: labels are weak, inferred "
        "from free text rather than direct image review, and the class "
        "distribution mirrors clinical prevalence, so normal studies overwhelm every "
        "pathological category. CheXNet demonstrated that a densely connected network could "
        "nonetheless attain competitive agreement with practising radiologists on pneumonia "
        "(Rajpurkar et al., 2017), while CheXpert addressed labelling ambiguity by encoding "
        "uncertainty explicitly (Irvin et al., 2019). Kumar, Grewal and Srivastava (2017) "
        "approached the same corpus through cascaded classifiers exploiting label "
        "co-occurrence, reporting that loss function selection materially influenced "
        "performance under imbalance."
    ),
    12: (
        "Whether to train from scratch or to transfer representations has attracted parallel "
        "attention. Baltruschat et al. (2019) conducted the most systematic comparison, "
        "evaluating ResNet variants with and without fine-tuning against networks trained "
        "exclusively on radiographs, and concluded that an "
        "X-ray-specific architecture incorporating non-image metadata outperformed ImageNet "
        "initialisation; their results indicate that frozen ImageNet features are adequate "
        "but not optimal for radiographic texture, a caveat directly relevant to the present "
        "design. Agughasi (2024) reached a partly contrasting conclusion for chronic "
        "obstructive pulmonary disease, finding that fine-tuned ResNet50 and VGG19 converged "
        "faster and more stably than a bespoke network, and that grayscale inputs outperformed "
        "colour. Brima et al. (2021) obtained 94% accuracy on a four-class pneumonia task "
        "using ResNet50 with conventional augmentation, and Fagbola and Success (2024) "
        "reported that generalisability varies substantially between pretrained backbones. "
        "Transfer learning is therefore a "
        "defensible default whose benefit remains task-dependent."
    ),
    13: (
        "Augmentation research divides along a clear methodological fault line. Shorten and "
        "Khoshgoftaar (2019) note that classical geometric and photometric transformations "
        "confer invariance rather than new information, and generative approaches were "
        "introduced to overcome that limit. Frid-Adar et al. (2018) demonstrated marked "
        "sensitivity and specificity gains for liver lesion classification using "
        "GAN-synthesised CT patches, and "
        "Salehinejad et al. (2018) extended the principle to chest radiography, reporting "
        "improved five-class classification when a deep convolutional GAN preferentially "
        "augmented sparse categories. During the COVID-19 pandemic the approach "
        "proliferated: Waheed et al. (2020) raised binary detection accuracy from 85% to 95% "
        "with an auxiliary classifier GAN, Karbhari et al. (2021) generated synthetic "
        "radiographs for a comparable two-class task, and Motamed, Rogalla and Khalvati "
        "(2021) proposed an inception-augmentation GAN for semi-supervised anomaly detection. "
        "Narin, Kaya and Pamuk (2021) meanwhile achieved very high accuracy on small binary "
        "COVID-19 cohorts using pretrained networks alone, suggesting the tasks on which "
        "generative augmentation was validated were already near saturation."
    ),
    14: (
        "A more sceptical literature has since emerged, and it motivates the present "
        "investigation. Sundaram and Hulkund (2021) evaluated GAN augmentation on CheXpert "
        "across several data regimes and found meaningful gains only when training data were "
        "severely restricted; at full size the benefit was minimal or negative. Segal et al. "
        "(2021) established through blinded radiologist assessment that progressively grown "
        "GAN outputs, though often mistaken for real images, fall short of clinical realism, "
        "particularly for fine-grained findings. Fedoruk et al. (2023) found StyleGAN2-ADA "
        "augmentation subpar against "
        "classical transformations for multi-class COVID-19 radiographs, and Schaudt et al. "
        "(2023) reported that image fidelity and downstream classification benefit were "
        "negatively correlated. Ali, Grönlund and Shah (2023) reviewed forty-three GAN studies "
        "in this domain and documented pervasive data bias, absent source code and negligible "
        "expert validation. Holste et al. (2022) independently showed that augmentation "
        "strategies such as MixUp fail on the rarest thoracic classes, where re-weighting and "
        "classifier re-training prove more effective."
    ),
    15: (
        "Two gaps therefore remain: positive GAN findings cluster in small, low-class-count "
        "settings, and negative findings seldom carry formal significance testing on a large "
        "multi-class corpus. This study addresses both."
    ),
    18: (
        "The investigation adopts a fully crossed factorial design in which augmentation "
        "regime and network architecture are manipulated while every other element of the "
        "pipeline is held constant. Four augmentation conditions were combined with three "
        "architectures to yield twelve independently trained models. Holding the data partition, "
        "resolution, batch size, optimiser, callbacks and evaluation protocol invariant ensures "
        "that any performance difference is attributable to the manipulated factors. A fixed random seed of 42 governed partitioning, shuffling and "
        "weight initialisation, a reproducibility requirement stressed by Lohani et al. (2025) "
        "and absent from much of the literature reviewed by Ali, Grönlund and Shah (2023)."
    ),
    20: (
        "A seven-class subset of frontal chest radiographs was assembled comprising No "
        "Finding, Infiltration, Atelectasis, Effusion, Nodule, Pneumothorax and Mass. "
        "Restricting the label space to single-finding studies converts the underlying "
        "multi-label problem into a tractable multi-class one, at the acknowledged cost of "
        "discarding co-occurrence information that Kumar, Grewal and Srivastava (2017) showed "
        "to be informative. The subset contains 85,116 images and is markedly imbalanced: No "
        "Finding accounts for 60,361 images, or 70.9%, whereas Mass, the rarest category, "
        "contributes 2,139, an imbalance ratio approaching 28:1. Stratified random sampling "
        "divided the images into 68,092 training and 17,024 validation cases in an 80/20 "
        "ratio, preserving class proportions in both partitions."
    ),
    22: (
        "All images were decoded as single-channel grayscale, resized to 224 x 224 pixels "
        "and rescaled to the unit interval. That resolution matches the native input "
        "geometry of both pretrained backbones and thereby avoids the arbitrary rescaling "
        "criticised by Ali, Grönlund and Shah (2023). Two normalisation paths were then applied: "
        "for the scratch-trained network, per-image standardisation removed "
        "exposure and contrast variation between acquisitions. For the transfer learning "
        "models, the single channel was tiled three times to satisfy the three-channel input "
        "requirement, then passed through the preprocessing function distributed with each "
        "pretrained architecture so that channel statistics matched "
        "those under which the ImageNet weights were learned. Retaining grayscale rather than "
        "artificial colourisation is consistent with Agughasi (2024)."
    ),
    24: (
        "The baseline condition applied no augmentation and establishes the reference against "
        "which the remaining conditions are judged. The classical condition applied random "
        "rotation of up to 0.05 of a turn, translation and zoom of up to 10%, and "
        "random horizontal flipping, with outputs clipped to the valid intensity range. These "
        "limits were kept conservative because aggressive geometric distortion risks "
        "displacing localised findings such as nodules beyond plausible positions. The "
        "intensity condition applied random brightness shifts of up to 0.10, contrast scaling "
        "between 0.80 and 1.20, and additive Gaussian noise with a standard deviation of 0.02, "
        "simulating the exposure and detector variability of different acquisition equipment "
        "rather than altering geometry."
    ),
    25: (
        "The generative condition employed a label-conditioned GAN following the conditional "
        "formulation of Mirza and Osindero (2014). The generator accepts a 100-dimensional "
        "latent vector concatenated with a 50-dimensional learned class embedding, projects this "
        "to a 7 x 7 x 256 feature volume, and upsamples through five transposed convolutional "
        "stages with batch normalisation and LeakyReLU activation to a 224 x 224 "
        "single-channel output under a hyperbolic tangent nonlinearity. The discriminator "
        "embeds the class label, projects it to a full-resolution plane, concatenates it with "
        "the input image and applies three strided convolutional blocks with LeakyReLU and "
        "dropout. Both were optimised with Adam at a learning rate of 1e-4 and a first moment "
        "decay of 0.5, the configuration conventionally used to stabilise adversarial "
        "training. Following training, 500 synthetic images were sampled "
        "per class, giving 3,500 images that were interleaved with the real training set "
        "before shuffling rather than replacing it, mirroring Salehinejad et al. (2018) and "
        "Frid-Adar et al. (2018)."
    ),
    27: (
        "The custom network comprises four convolutional blocks with 32, 64, 128 and 256 "
        "filters, each combining a 3 x 3 convolution, batch normalisation, rectified linear "
        "activation, max pooling and increasing dropout. Global average pooling replaces "
        "flattening to constrain parameter count, and a 256-unit dense layer with batch "
        "normalisation and dropout precedes the softmax output. ResNet50 (He et al., 2016) and "
        "VGG19 (Simonyan and Zisserman, 2015) were instantiated with ImageNet weights (Deng et "
        "al., 2009) and their convolutional bases frozen, with an identical head of global "
        "average pooling, a 256-unit rectified linear layer, dropout at 0.30 and a softmax "
        "output. Freezing the backbone isolates the "
        "effect of augmentation on the trainable head, consistent with Fagbola and Success "
        "(2024), although Baltruschat et al. (2019) indicate that fine-tuning would likely "
        "raise absolute performance."
    ),
    29: (
        "All classifiers minimised sparse categorical cross-entropy using Adam (Kingma and "
        "Ba, 2015) with a batch size of 32, applying a learning rate of 3e-4 for the "
        "scratch-trained network and 1e-3 for the transfer learning heads. Class imbalance "
        "was addressed through square-root-balanced class weights, the square root of the "
        "conventional inverse-frequency weights renormalised to unit mean. Full "
        "inverse-frequency weighting was rejected because a 28:1 ratio produces gradient "
        "contributions from rare classes large enough to destabilise optimisation. "
        "Training ran for up to twelve epochs under three callbacks monitoring validation loss: "
        "checkpointing retained the best weights, early stopping with a patience of four "
        "epochs terminated unproductive runs and restored the optimal parameters, and learning "
        "rate reduction halved the step size after two stagnant epochs down to a floor of "
        "1e-6. Selecting on validation loss rather than accuracy avoids rewarding the "
        "degenerate majority-class solution."
    ),
    31: (
        "Because accuracy is uninformative when a single class occupies 70.9% of the data, "
        "performance is reported primarily as macro-averaged precision, recall and F1, which "
        "weight every class equally regardless of prevalence. Accuracy is retained for "
        "comparability against the 0.709 majority-class prior. "
        "Reliability was assessed by non-parametric bootstrap resampling of each confusion "
        "matrix, drawing 2,000 multinomial replicates at the observed cell probabilities with "
        "a seed of 42 to obtain 95% percentile intervals for macro-F1. Pairwise contrasts "
        "were evaluated on the distribution of paired differences, with two-sided p-values from "
        "the proportion of replicates in which the difference reversed sign and significance "
        "declared at 0.05. This follows standard bootstrap practice (Efron and Tibshirani, "
        "1993) and requires no distributional assumption. The recorded per-epoch histories "
        "were additionally inspected to establish whether the reported figures reflect a "
        "well-conditioned fit."
    ),
    32: (
        "Figure 1. Experimental workflow for the twelve augmentation-by-architecture "
        "experiments, from metadata filtering and stratified partitioning through the four "
        "augmentation branches, model construction and bootstrap significance testing."
    ),
    35: (
        "Table 1 reports the four evaluation metrics for all twelve experiments on the "
        "held-out validation set of 17,024 images. The most consequential observation is that "
        "no configuration surpassed the majority-class prior of 0.709: the highest accuracy "
        "recorded, 0.707 for the VGG19 baseline, falls marginally below the value obtainable "
        "by predicting No Finding for every study. Accuracy is therefore uninformative as a "
        "discriminator between conditions, and the analysis proceeds on macro-averaged "
        "measures."
    ),
    36: (
        "Macro-F1 ranged from 0.038 to 0.175. The strongest configuration was VGG19 with "
        "intensity augmentation at 0.175, followed by the ResNet50 baseline at 0.174. "
        "Conditional GAN augmentation occupied the lowest rank within every architecture, at "
        "0.038 for the custom network, 0.135 for ResNet50 and 0.127 for VGG19. Macro-precision "
        "varied more widely, reaching 0.455 for ResNet50 with classical augmentation, but this "
        "reflects extreme conservatism rather than skill: that model assigned only 183 of "
        "17,024 studies to any class other than No Finding, so its few minority predictions "
        "were disproportionately correct while recall fell to 0.149."
    ),
    38: (
        "Bootstrap resampling with 2,000 replicates yielded narrow intervals, reflecting the "
        "substantial validation sample. The best configuration, VGG19 with intensity "
        "augmentation, had a 95% confidence interval of 0.169 to 0.181, while the ResNet50 "
        "baseline spanned 0.168 to 0.180. These intervals overlap almost entirely, and the "
        "paired contrast confirms that the 0.0015 difference between them is not significant "
        "(p = 0.753), so any claim that VGG19 with intensity augmentation is the superior model "
        "would overstate the data. By contrast, the interval for the GAN-augmented custom "
        "network, 0.035 to 0.042, is disjoint from every other interval by a wide margin."
    ),
    40: (
        "Holding augmentation constant and contrasting architectures produced a less "
        "consistent ordering than might be expected. Under the baseline regime, ResNet50 "
        "significantly outperformed both the custom network, by 0.021, and VGG19, by 0.038, at "
        "p < 0.001, yet that advantage did not survive augmentation. Under classical "
        "augmentation the custom network and VGG19 were indistinguishable (p = 0.745), and "
        "under intensity augmentation the difference between the custom network and ResNet50 "
        "likewise failed to reach significance (p = 0.074). The merit of an architecture is "
        "therefore conditional on the augmentation applied to it."
    ),
    42: (
        "The augmentation effect proved strongly architecture-dependent, which is the central "
        "analytical finding. For the scratch-trained network, classical and intensity "
        "augmentation produced modest but significant gains over the baseline, of 0.011 "
        "(p = 0.009) and 0.009 (p = 0.034); for VGG19 the corresponding gains were larger, "
        "0.027 and 0.040, both at p < 0.001. For ResNet50, by contrast, every regime "
        "significantly degraded macro-F1 relative to its own baseline, by 0.043 for classical, "
        "0.020 for intensity and 0.035 for the GAN condition, all at p < 0.001. No single "
        "policy can therefore be recommended across architectures, and the frozen ResNet50 "
        "features appear least tolerant of input perturbation."
    ),
    43: (
        "Conditional GAN augmentation was harmful without exception. Relative to their "
        "respective baselines it reduced macro-F1 by 0.114 for the custom network "
        "(p < 0.001), by 0.035 for ResNet50 (p < 0.001) and by 0.009 for VGG19 (p = 0.001). "
        "The ResNet50 contrast used a confusion matrix regenerated from the saved checkpoint, "
        "yielding a macro-F1 of 0.139 against the 0.135 in Table 1; the discrepancy is "
        "immaterial, since both indicate significant deterioration."
    ),
    44: "4.6 Failure modes and confusion structure",
    45: (
        "Per-class analysis clarifies the mechanism (Figures 4 to 6). The GAN-augmented "
        "custom network did not merely underperform but inverted: it assigned 12,190 of "
        "17,024 studies to Effusion, a class representing 791 true cases, achieving 0.880 "
        "recall there while collapsing to 0.010 recall for No Finding (Figure 6). Accuracy of "
        "0.058 indicates a systematic mapping onto a minority class rather than random "
        "guessing. The GAN-augmented VGG19 showed the opposite pathology, producing only No "
        "Finding and Infiltration predictions for 17,018 of 17,024 studies."
    ),
    46: (
        "Degeneracy toward the majority class was pervasive, and Figures 4 and 5 show its form "
        "precisely. Both direct almost all "
        "predictions into the No Finding column, 15,434 of 17,024 for VGG19 with intensity "
        "augmentation and 16,013 for the ResNet50 baseline, so each minority row loses most of "
        "its cases to it: 88.5% of Infiltration, 85.8% of Atelectasis, 94.1% of Nodule and 92.1% "
        "of Mass under VGG19, and between 91.1% and 97.0% of the same four classes under "
        "ResNet50. Effusion is the only finding partially "
        "recovered, at an F1 of 0.228 and 0.297, and doubles as the principal secondary sink, "
        "absorbing 431 No Finding studies under VGG19. Nodule and Mass returned an F1 of "
        "exactly zero in nine of the twelve configurations and exceeded 0.055 in none, against "
        "No Finding F1 of 0.820 and 0.827 for the two leaders. Square-root class weighting was "
        "therefore insufficient against an imbalance ratio approaching 28:1 combined with "
        "frozen feature extractors."
    ),
    50: (
        "The evidence assembled here aligns with the cautious strand of the augmentation "
        "literature rather than the optimistic reports that dominate work on synthetic "
        "radiographs. The most instructive comparison is with Sundaram and Hulkund (2021), whose "
        "GAN augmentation improved CheXpert classification substantially at one per cent of the "
        "corpus but delivered minimal or negative benefit at full scale. The present "
        "experiments ran at full scale "
        "throughout, with 68,092 real training images against 3,500 synthetic additions, and "
        "returned uniformly negative outcomes; the value of synthetic augmentation therefore "
        "appears contingent on data scarcity rather than intrinsic. That also explains the "
        "apparent conflict with Waheed et al. (2020) and Karbhari et al. (2021), whose gains "
        "came from two-category COVID-19 cohorts of a few thousand images, and with Frid-Adar "
        "et al. (2018), who succeeded with small lesion patches: in each case a generator need "
        "only capture coarse distinguishing features, a far less demanding task than "
        "synthesising full radiographs across seven classes."
    ),
    51: (
        "The mechanism underlying the harm merits scrutiny. Segal et al. (2021) showed that even "
        "progressively grown GANs trained to convergence produce images radiologists can "
        "distinguish from real studies, and a conditional generator given a limited epoch "
        "budget at full resolution cannot plausibly exceed that standard. The synthetic images "
        "are better "
        "characterised as mislabelled inputs than as augmentation, which accounts for the "
        "pattern of damage observed: the scratch-trained network, whose features were still "
        "being learned, was catastrophically affected, whereas the transfer learning heads, "
        "optimising only a shallow classifier over fixed representations, absorbed the "
        "corruption with far less disruption. Schaudt et al. (2023) reached a convergent "
        "conclusion, reporting a negative correlation between generated fidelity and "
        "downstream gain, and Fedoruk et al. (2023) likewise found StyleGAN2-ADA augmentation "
        "inferior to classical transformation. Ali, Grönlund and Shah (2023) attribute much of "
        "the field's optimism to data bias and absent independent validation."
    ),
    52: (
        "The architecture-dependence of the augmentation effect deserves equal emphasis, "
        "since it is rarely reported. That identical classical and intensity policies "
        "improved VGG19 and the custom network while significantly degrading ResNet50 "
        "indicates an interaction between augmentation and representation that "
        "single-architecture studies cannot detect. A plausible explanation lies in the batch "
        "normalisation statistics embedded in the frozen ResNet50 backbone, estimated on "
        "unperturbed ImageNet data and unable to adapt when input distributions shift. "
        "Baltruschat et al. (2019) anticipated the point in showing that "
        "frozen ImageNet features are suboptimal for radiographic texture, and Agughasi "
        "(2024) reported that fine-tuning confers stability that frozen extraction does not."
    ),
    53: (
        "Several limitations qualify these conclusions. A single stratified split was used, so "
        "the bootstrap intervals capture validation sampling variability but not run-to-run "
        "variance. Validation data served for both model selection and reporting, which risks "
        "mild optimism, and no external cohort was available. The convolutional bases remained "
        "frozen and the GAN received a modest training budget, so the study establishes that "
        "this configuration fails rather than that conditional GANs cannot succeed. Generated "
        "image quality was not quantified using Fréchet Inception Distance, which would have "
        "addressed fidelity directly. Finally, collapsing a multi-label problem to single "
        "findings discards the co-occurrence structure that Kumar, Grewal and Srivastava "
        "(2017) exploited."
    ),
    55: (
        "This study evaluated conditional GAN augmentation against classical geometric and "
        "intensity-based alternatives for seven-class chest radiograph classification, using "
        "three architectures and bootstrap significance testing across twelve controlled "
        "experiments. The central finding is unambiguous: label-conditioned synthetic "
        "augmentation did not improve macro-averaged F1 for any architecture and significantly "
        "degraded it in every comparison, culminating in predictive collapse for the network "
        "trained from scratch. Classical and intensity augmentation produced small but "
        "reliable gains for VGG19 and the custom network yet significantly harmed the frozen "
        "ResNet50, showing that augmentation policy interacts with architecture. The best "
        "macro-F1 achieved, 0.175 for VGG19 with intensity "
        "augmentation, was statistically indistinguishable from the ResNet50 baseline, and no "
        "configuration exceeded the 0.709 majority-class prior in accuracy."
    ),
    56: (
        "The take-home message is twofold. First, synthetic augmentation is not a general "
        "remedy for class imbalance; where real data are plentiful and a generator is not "
        "trained to convergence, its output functions as label noise and degrades "
        "minority-class recovery. Second, severe imbalance in multi-class radiography is not "
        "solved by augmentation of any variety, and the near-total failure on Nodule and Mass "
        "indicates that loss reweighting, classifier re-training or fine-tuning the feature "
        "extractor are more promising. Future work should fine-tune the convolutional bases, "
        "quantify synthetic fidelity, and repeat the comparison across multiple seeds and an "
        "external cohort."
    ),
}

# ------------------------------------------------------------------- additions

OVERFIT_HEADING = "4.5 Training dynamics and evidence of overfitting"

OVERFIT_TEXT = (
    "Because a macro-F1 difference of this magnitude could reflect overfitting as readily as "
    "representational advantage, the recorded per-epoch histories of the two leading "
    "configurations were examined directly (Figures 2 and 3). Neither shows the canonical "
    "overfitting signature of training accuracy rising above validation accuracy: training "
    "accuracy peaked at 0.546 for VGG19 with intensity augmentation, below its terminal "
    "validation accuracy of 0.555, and at 0.633 for the ResNet50 baseline against 0.669. "
    "Divergence is nevertheless visible in the VGG19 loss curves, where validation loss "
    "reached its minimum of 1.178 at epoch 2 and then rose to 1.694 by epoch 4 while training "
    "loss continued to fall. The ResNet50 baseline shows no such trend: its validation loss "
    "minimum of 1.057 fell at epoch 10 of 12, below the first-epoch value of 1.254. The "
    "dominant feature of both histories is instability rather than memorisation, validation "
    "accuracy spanning 0.179 to 0.679 for VGG19 and 0.468 to 0.700 for ResNet50. Because "
    "early stopping restored the minimum-validation-loss epoch in each case, epoch 2 and "
    "epoch 10, training was halted at rather than beyond the optimum, and the low ceiling on "
    "training accuracy indicates underfitting of the minority classes rather than excessive "
    "fit to the majority."
)

CAP_FIG2 = (
    "Figure 2. Training and validation accuracy per epoch, from the recorded histories, for "
    "(a) VGG19 with intensity augmentation and (b) the ResNet50 baseline. Dashed lines mark "
    "the minimum-validation-loss epoch whose restored weights supplied the reported metrics."
)

CAP_FIG3 = (
    "Figure 3. Training and validation cross-entropy loss for the same two configurations. "
    "Validation loss rises after epoch 2 in panel (a) while training loss falls; in panel (b) "
    "it reaches its minimum at epoch 10 of 12."
)

CAP_FIG4 = (
    "Figure 4. Confusion matrix for VGG19 with intensity augmentation, the highest-ranking "
    "configuration by macro-F1, on the 17,024-image validation set. Annotations are counts "
    "and shading is normalised within each row, so the diagonal reads as per-class recall."
)

CAP_FIG5 = (
    "Figure 5. Confusion matrix for the ResNet50 baseline, whose macro-F1 of 0.174 is "
    "statistically indistinguishable from that of Figure 4. Annotation and shading "
    "conventions follow Figure 4."
)

CAP_FIG6 = (
    "Figure 6. Confusion matrix for the custom CNN with conditional GAN augmentation, retained "
    "as a contrastive failure case in which predictions collapse onto the Effusion column. No "
    "learning curves are shown because this run's recorded history could not be matched to the "
    "evaluated checkpoint."
)

DISCUSSION_NEW = (
    "The training histories bear on how the ceiling reported here should be read. The absence "
    "of a memorisation signature means it cannot be attributed to overfitting and will not be "
    "lifted by stronger regularisation; the binding constraints are the frozen representations "
    "and the imbalance itself, consistent with Baltruschat et al. (2019). The volatility of "
    "the validation curves, most pronounced in the VGG19 run whose validation accuracy fell to "
    "0.179 before recovering, also means that selection on a single split rewards a favourable "
    "epoch as much as a favourable model, reinforcing the indistinguishability of the two "
    "leading configurations. The matrices locate the residual difficulty in the tail, where "
    "Nodule and Mass are effectively never predicted, the pattern Holste et al. (2022) "
    "addressed through re-weighting and classifier re-training rather than augmentation."
)

# ------------------------------------------------------------------- utilities


def set_text(paragraph, text):
    runs = paragraph.runs
    if not runs:
        paragraph.add_run(text)
        return
    runs[0].text = text
    for r in runs[1:]:
        r._element.getparent().remove(r._element)


def _clone_formatting(new_p, model_p):
    if new_p._p.pPr is not None:
        new_p._p.remove(new_p._p.pPr)
    if model_p._p.pPr is not None:
        new_p._p.insert(0, copy.deepcopy(model_p._p.pPr))
    if model_p.runs and new_p.runs:
        src_rpr = model_p.runs[0]._element.rPr
        if src_rpr is not None:
            dst = new_p.runs[0]._element
            if dst.rPr is not None:
                dst.remove(dst.rPr)
            dst.insert(0, copy.deepcopy(src_rpr))


def insert_text_before(ref_p, text, model_p):
    new_p = ref_p.insert_paragraph_before(text, style=model_p.style)
    _clone_formatting(new_p, model_p)
    return new_p


def insert_image_before(ref_p, image_path, width_in, style_source):
    new_p = ref_p.insert_paragraph_before("", style=style_source.style)
    new_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = new_p.paragraph_format
    pf.line_spacing = 1.0
    pf.space_before = Pt(8)
    pf.space_after = Pt(2)
    new_p.add_run().add_picture(image_path, width=Inches(width_in))
    return new_p


def body_word_count(doc):
    """Words from the Abstract heading to the end of the Conclusion."""
    texts = [p.text for p in doc.paragraphs]
    start = texts.index("Abstract")
    end = texts.index("References")
    body = texts[start:end]
    total = sum(len(t.split()) for t in body)
    captions = sum(
        len(t.split())
        for t in body
        if t.startswith("Figure ") or t.startswith("Table ")
    )
    return total, captions


# ------------------------------------------------------------------------ main


def main(dry):
    doc = Document(DOCX)
    paras = list(doc.paragraphs)

    for idx, text in REPLACE.items():
        set_text(paras[idx], text)

    body_model = paras[35]      # justified body text
    label_model = paras[44]     # bold subsection label
    caption_model = paras[47]   # italic caption

    # Figure 1: embed the workflow diagram above its existing caption.
    insert_image_before(paras[32], WORKFLOW, 6.0, body_model)
    paras[32].alignment = None
    _clone_formatting(paras[32], caption_model)

    # New Section 4.5 (training dynamics) immediately before the failure-modes label.
    anchor = paras[44]
    insert_text_before(anchor, OVERFIT_HEADING, label_model)
    insert_text_before(anchor, OVERFIT_TEXT, body_model)
    insert_image_before(anchor, os.path.join(FIG, "fig2_accuracy_curves.png"), 6.0, body_model)
    insert_text_before(anchor, CAP_FIG2, caption_model)
    insert_image_before(anchor, os.path.join(FIG, "fig3_loss_curves.png"), 6.0, body_model)
    insert_text_before(anchor, CAP_FIG3, caption_model)

    # Confusion-matrix figures replace the old combined Figure 2 caption after Table 1.
    old_cap = paras[48]
    for image, caption in [
        ("fig4_cm_intensity_vgg19.png", CAP_FIG4),
        ("fig5_cm_baseline_resnet50.png", CAP_FIG5),
        ("fig6_cm_gan_cnn.png", CAP_FIG6),
    ]:
        insert_image_before(old_cap, os.path.join(FIG, image), 5.3, body_model)
        insert_text_before(old_cap, caption, caption_model)
    old_cap._p.getparent().remove(old_cap._p)

    # New discussion paragraph linking curves and matrices to interpretation.
    insert_text_before(paras[53], DISCUSSION_NEW, body_model)

    total, captions = body_word_count(doc)
    print(f"body words (Abstract..Conclusion): {total}")
    print(f"  of which figure/table captions:  {captions}")
    print(f"  excluding captions:              {total - captions}")

    if dry:
        print("dry run: not written")
        return

    backup = DOCX.replace(".docx", "_prefigures_backup.docx")
    if not os.path.exists(backup):
        shutil.copy2(DOCX, backup)
        print("backup:", backup)
    doc.save(DOCX)
    print("saved:", DOCX)


if __name__ == "__main__":
    main("--dry" in sys.argv)
